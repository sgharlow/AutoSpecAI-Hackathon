"""
AutoSpec.AI Process Lambda Function with Intelligent Caching

Enhanced version of the process function that demonstrates comprehensive
caching strategies for performance optimization.

Key caching optimizations:
- Document content caching to avoid repeated S3 downloads
- Bedrock AI analysis caching based on document content hash
- Configuration caching for faster function initialization
- Template caching for consistent output formatting
"""

import json
import boto3
import uuid
import time
import os
import hashlib
import asyncio
from datetime import datetime, timezone
from typing import Dict, Any, Optional, List
from concurrent.futures import ThreadPoolExecutor
import contextlib

# Import our enhanced shared modules
import sys
sys.path.append('/opt/python/lib/python3.11/site-packages')
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from shared import (
    ErrorHandler, ExternalServiceError, DatabaseError,
    get_config, get_logger, LogContext, log_performance,
    retry_with_backoff, retry_bedrock_call, RetryStrategy,
    CircuitBreaker, HealthChecker, AWSHealthChecks
)

# Import caching infrastructure
from shared.cache_manager import get_cache_manager, cached

# Initialize configuration and utilities with caching
cache_manager = get_cache_manager()
config = get_config()
error_handler = ErrorHandler('process', config.environment)
logger = get_logger('process', 'process', config.environment)

# Initialize optimized AWS clients
dynamodb = boto3.resource('dynamodb', config=boto3.session.Config(
    retries={'max_attempts': 5, 'mode': 'adaptive'},
    max_pool_connections=50
))

bedrock_client = boto3.client('bedrock-runtime', config=boto3.session.Config(
    retries={'max_attempts': 3, 'mode': 'adaptive'},
    read_timeout=300,
    connect_timeout=60
))

s3_client = boto3.client('s3', config=boto3.session.Config(
    retries={'max_attempts': 3, 'mode': 'adaptive'}
))

# Environment variables
DOCUMENT_BUCKET = os.environ.get('DOCUMENT_BUCKET')
PROCESSING_BUCKET = os.environ.get('PROCESSING_BUCKET', DOCUMENT_BUCKET)
HISTORY_TABLE = os.environ.get('HISTORY_TABLE')
CACHE_TABLE = os.environ.get('CACHE_TABLE', 'autospec-ai-cache')

# Performance optimization constants
MAX_WORKERS = 3
BATCH_SIZE = 10
TEXT_CHUNK_SIZE = 7000

# Circuit breaker for Bedrock
bedrock_circuit_breaker = CircuitBreaker.get_instance('bedrock', failure_threshold=5, recovery_timeout_seconds=120)

@error_handler.lambda_handler
def handler(event, context):
    """
    Optimized Lambda handler with comprehensive caching strategies.
    """
    start_time = time.time()
    
    with LogContext(
        aws_request_id=context.aws_request_id,
        function_name=context.function_name,
        memory_limit_mb=context.memory_limit_in_mb
    ):
        logger.info("Process function started with caching",
                   event_records=len(event.get('Records', [])),
                   memory_limit=context.memory_limit_in_mb,
                   cache_stats=cache_manager.get_stats())
        
        # Health check if no records
        if not event.get('Records'):
            return handle_health_check()
        
        # Process S3 events efficiently with caching
        results = []
        with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
            futures = []
            
            for record in event['Records']:
                if record.get('eventSource') == 'aws:s3':
                    future = executor.submit(process_s3_record_with_caching, record)
                    futures.append(future)
            
            # Collect results
            for future in futures:
                try:
                    result = future.result(timeout=600)
                    results.append(result)
                except Exception as e:
                    logger.error("Failed to process S3 record", error=str(e))
                    results.append({'status': 'failed', 'error': str(e)})
        
        duration = time.time() - start_time
        final_cache_stats = cache_manager.get_stats()
        
        logger.info("Process function completed with caching",
                   processed_records=len(results),
                   successful_records=len([r for r in results if r.get('status') == 'success']),
                   duration_seconds=duration,
                   cache_stats=final_cache_stats)
        
        return {
            'statusCode': 200,
            'body': json.dumps({
                'processed_records': len(results),
                'results': results,
                'cache_stats': final_cache_stats
            })
        }

@log_performance('s3_record_processing_cached')
def process_s3_record_with_caching(record: Dict[str, Any]) -> Dict[str, Any]:
    """Process a single S3 record with comprehensive caching."""
    
    try:
        # Extract S3 information
        s3_info = record['s3']
        bucket_name = s3_info['bucket']['name']
        object_key = s3_info['object']['key']
        
        logger.info("Processing S3 object with caching",
                   bucket=bucket_name,
                   key=object_key,
                   size_bytes=s3_info['object'].get('size'))
        
        # Use cached GSI lookup for DynamoDB record
        db_record = find_record_by_s3_key_cached(object_key)
        if not db_record:
            raise DatabaseError(f"No DynamoDB record found for S3 key: {object_key}")
        
        request_id = db_record['requestId']
        
        with LogContext(request_id=request_id):
            # Update processing stage
            update_processing_stage_optimized(request_id, 'ai_processing_in_progress')
            
            # Download and process document with caching
            document_content, document_hash = download_document_with_caching(bucket_name, object_key)
            
            # Check if we have cached analysis for this document content
            cached_analysis = cache_manager.get_cached_analysis(request_id, document_hash)
            
            if cached_analysis:
                logger.info("Using cached AI analysis", 
                           request_id=request_id, 
                           document_hash=document_hash[:8])
                ai_analysis = cached_analysis
            else:
                # Extract text with caching
                text_content = extract_text_with_caching(document_content, object_key, document_hash)
                
                # Process with Bedrock and cache result
                ai_analysis = process_with_bedrock_cached(text_content, db_record, document_hash)
                
                # Cache the analysis result
                cache_manager.set_cached_analysis(request_id, document_hash, ai_analysis)
            
            # Store results efficiently
            store_results_optimized(request_id, ai_analysis, db_record)
            
            # Update completion status
            update_processing_stage_optimized(request_id, 'ai_processing_complete', {
                'ai_analysis_length': len(json.dumps(ai_analysis)),
                'processing_duration_seconds': time.time(),
                'cache_hit': cached_analysis is not None,
                'document_hash': document_hash
            })
            
            logger.info("S3 record processed successfully with caching", 
                       request_id=request_id,
                       used_cache=cached_analysis is not None)
            
            return {
                'status': 'success',
                'request_id': request_id,
                'object_key': object_key,
                'cache_hit': cached_analysis is not None,
                'document_hash': document_hash
            }
            
    except Exception as e:
        logger.error("S3 record processing failed",
                    bucket=bucket_name,
                    key=object_key,
                    error=str(e))
        
        # Try to update error status if we have request_id
        if 'request_id' in locals():
            try:
                update_processing_stage_optimized(request_id, 'failed', {
                    'error_message': str(e),
                    'error_type': type(e).__name__
                })
            except Exception as update_error:
                logger.error("Failed to update error status", error=str(update_error))
        
        return {
            'status': 'failed',
            'error': str(e),
            'object_key': object_key
        }

@cached(cache_type='document_meta', ttl=1800)
def find_record_by_s3_key_cached(s3_key: str) -> Optional[Dict[str, Any]]:
    """Find DynamoDB record using S3KeyIndex GSI with caching."""
    
    try:
        table = dynamodb.Table(HISTORY_TABLE)
        
        # Use GSI for efficient query
        response = table.query(
            IndexName='S3KeyIndex',
            KeyConditionExpression='s3Key = :s3_key',
            ExpressionAttributeValues={':s3_key': s3_key},
            Limit=1,
            ScanIndexForward=False
        )
        
        if response['Items']:
            logger.info("Found record using cached S3KeyIndex lookup",
                       s3_key=s3_key,
                       consumed_capacity=response.get('ConsumedCapacity'))
            return response['Items'][0]
        else:
            logger.warning("No record found for S3 key", s3_key=s3_key)
            return None
            
    except Exception as e:
        logger.error("Failed to query S3KeyIndex", s3_key=s3_key, error=str(e))
        raise DatabaseError(f"Failed to find record for S3 key {s3_key}: {str(e)}")

def download_document_with_caching(bucket_name: str, object_key: str) -> tuple[bytes, str]:
    """Download document with content-based caching."""
    
    # Generate cache key based on S3 object metadata
    cache_key = f"document:{bucket_name}:{object_key}"
    
    # Try to get from cache first
    cached_result = cache_manager.get(cache_key, 'bedrock_result')
    if cached_result:
        logger.info("Using cached document content", 
                   bucket=bucket_name, 
                   key=object_key)
        return cached_result['content'], cached_result['hash']
    
    try:
        # Download from S3
        response = s3_client.get_object(Bucket=bucket_name, Key=object_key)
        
        content_length = response['ContentLength']
        logger.info("Downloading document from S3",
                   bucket=bucket_name,
                   key=object_key,
                   size_bytes=content_length)
        
        # Read content efficiently
        body = response['Body']
        chunks = []
        bytes_read = 0
        
        while True:
            chunk = body.read(8192)  # 8KB chunks
            if not chunk:
                break
            chunks.append(chunk)
            bytes_read += len(chunk)
        
        content = b''.join(chunks)
        
        # Calculate content hash for deduplication
        content_hash = hashlib.sha256(content).hexdigest()
        
        # Cache the document content (but only if it's not too large)
        if len(content) < 10 * 1024 * 1024:  # 10MB limit for caching
            cache_data = {
                'content': content,
                'hash': content_hash,
                'size': len(content)
            }
            cache_manager.set(cache_key, cache_data, 'bedrock_result', ttl=3600)
        
        return content, content_hash
        
    except Exception as e:
        logger.error("Document download failed", 
                    bucket=bucket_name, 
                    key=object_key, 
                    error=str(e))
        raise ExternalServiceError(f"Failed to download document: {str(e)}")

def extract_text_with_caching(document_content: bytes, filename: str, document_hash: str) -> str:
    """Extract text with content-based caching."""
    
    # Generate cache key based on document hash and filename
    cache_key = f"text_extraction:{document_hash}:{os.path.splitext(filename)[1]}"
    
    # Try to get from cache first
    cached_text = cache_manager.get(cache_key, 'analysis')
    if cached_text:
        logger.info("Using cached text extraction", 
                   filename=filename, 
                   document_hash=document_hash[:8])
        return cached_text
    
    # Extract text using the original logic
    file_extension = os.path.splitext(filename)[1].lower()
    
    try:
        if file_extension == '.txt':
            # Efficient text decoding with fallback encodings
            for encoding in ['utf-8', 'utf-16', 'latin-1']:
                try:
                    text_content = document_content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                text_content = document_content.decode('utf-8', errors='ignore')
                
        elif file_extension == '.pdf':
            text_content = extract_pdf_text_optimized(document_content)
            
        elif file_extension in ['.docx', '.doc']:
            text_content = extract_docx_text_optimized(document_content)
            
        else:
            logger.warning("Unsupported file type for text extraction", 
                          extension=file_extension)
            text_content = f"[Unsupported file type: {file_extension}]"
        
        # Cache the extracted text
        cache_manager.set(cache_key, text_content, 'analysis', ttl=7200)  # 2 hours
        
        logger.info("Text extracted and cached", 
                   filename=filename,
                   text_length=len(text_content),
                   document_hash=document_hash[:8])
        
        return text_content
            
    except Exception as e:
        logger.error("Text extraction failed", 
                    filename=filename, 
                    error=str(e))
        return f"[Text extraction failed: {str(e)}]"

def extract_pdf_text_optimized(content: bytes) -> str:
    """Optimized PDF text extraction with memory management."""
    try:
        import PyPDF2
        import io
        
        pdf_stream = io.BytesIO(content)
        pdf_reader = PyPDF2.PdfReader(pdf_stream)
        
        text_parts = []
        total_pages = len(pdf_reader.pages)
        
        logger.info("Extracting PDF text", total_pages=total_pages)
        
        # Process pages in batches to manage memory
        batch_size = 10
        for i in range(0, total_pages, batch_size):
            batch_end = min(i + batch_size, total_pages)
            batch_text = []
            
            for page_num in range(i, batch_end):
                try:
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text.strip():
                        batch_text.append(page_text)
                except Exception as page_error:
                    logger.warning(f"Failed to extract page {page_num}", error=str(page_error))
            
            text_parts.extend(batch_text)
            
            # Log progress for large documents
            if total_pages > 50:
                logger.info("PDF extraction progress",
                           pages_processed=batch_end,
                           total_pages=total_pages,
                           progress_percent=round((batch_end / total_pages) * 100, 1))
        
        return '\n'.join(text_parts)
        
    except ImportError:
        logger.error("PyPDF2 not available for PDF processing")
        return "[PDF processing library not available]"
    except Exception as e:
        logger.error("PDF text extraction failed", error=str(e))
        return f"[PDF extraction failed: {str(e)}]"

def extract_docx_text_optimized(content: bytes) -> str:
    """Optimized DOCX text extraction."""
    try:
        from docx import Document
        import io
        
        doc_stream = io.BytesIO(content)
        doc = Document(doc_stream)
        
        text_parts = []
        total_paragraphs = len(doc.paragraphs)
        
        logger.info("Extracting DOCX text", total_paragraphs=total_paragraphs)
        
        # Process paragraphs efficiently
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:
                text_parts.append(text)
            
            # Log progress for large documents
            if total_paragraphs > 1000 and i % 500 == 0:
                logger.info("DOCX extraction progress",
                           paragraphs_processed=i,
                           total_paragraphs=total_paragraphs)
        
        return '\n'.join(text_parts)
        
    except ImportError:
        logger.error("python-docx not available for DOCX processing")
        return "[DOCX processing library not available]"
    except Exception as e:
        logger.error("DOCX text extraction failed", error=str(e))
        return f"[DOCX extraction failed: {str(e)}]"

def process_with_bedrock_cached(text_content: str, db_record: Dict[str, Any], document_hash: str) -> Dict[str, Any]:
    """Process document with Bedrock using intelligent caching."""
    
    # Truncate text if too long
    max_chars = config.processing.text_extraction.prompt_char_limit
    if len(text_content) > max_chars:
        logger.info("Truncating text for processing",
                   original_length=len(text_content),
                   truncated_length=max_chars)
        text_content = text_content[:max_chars] + "\n[... content truncated ...]"
    
    # Generate cache key based on content hash and processing parameters
    processing_params = {
        'model_id': config.bedrock.model_id,
        'max_tokens': config.bedrock.max_tokens,
        'temperature': config.bedrock.temperature,
        'text_hash': hashlib.sha256(text_content.encode()).hexdigest()[:16]
    }
    
    cache_key = f"bedrock_analysis:{document_hash}:{hashlib.md5(json.dumps(processing_params, sort_keys=True).encode()).hexdigest()[:8]}"
    
    # Try to get from cache first
    cached_result = cache_manager.get(cache_key, 'bedrock_result')
    if cached_result:
        logger.info("Using cached Bedrock analysis",
                   document_hash=document_hash[:8],
                   cache_key=cache_key[:16])
        return cached_result
    
    # Prepare optimized prompt with caching
    filename = db_record.get('filename', 'document')
    sender_email = db_record.get('senderEmail', 'unknown')
    
    prompt = get_cached_prompt_template(text_content, filename, sender_email)
    
    # Process with circuit breaker protection
    try:
        with bedrock_circuit_breaker:
            result = call_bedrock_with_retry(prompt, filename)
            
            # Cache the result
            cache_manager.set(cache_key, result, 'bedrock_result', ttl=86400)  # 24 hours
            
            logger.info("Bedrock analysis completed and cached",
                       document_hash=document_hash[:8],
                       result_size=len(json.dumps(result)))
            
            return result
            
    except Exception as e:
        logger.error("Bedrock processing failed", error=str(e))
        # Return fallback analysis
        return create_fallback_analysis(text_content, filename)

def get_cached_prompt_template(text_content: str, filename: str, sender_email: str) -> str:
    """Get cached prompt template with variable substitution."""
    
    template_vars = {
        'filename': filename,
        'sender_email': sender_email,
        'content_length': len(text_content),
        'text_content': text_content
    }
    
    # Try to get cached template
    cached_prompt = cache_manager.get_cached_template('bedrock_analysis', template_vars)
    if cached_prompt:
        return cached_prompt
    
    # Generate prompt
    prompt = create_optimized_prompt(text_content, filename, sender_email)
    
    # Cache the template
    cache_manager.set_cached_template('bedrock_analysis', template_vars, prompt)
    
    return prompt

def create_optimized_prompt(text_content: str, filename: str, sender_email: str) -> str:
    """Create an optimized prompt for Bedrock processing."""
    
    return f"""As an expert systems analyst, analyze the following document and extract comprehensive system requirements.

Document Information:
- Filename: {filename}
- Requester: {sender_email}
- Content Length: {len(text_content)} characters

Document Content:
{text_content}

Please provide a detailed analysis in the following JSON structure:
{{
  "document_summary": "Brief overview of the document purpose and scope",
  "functional_requirements": [
    {{
      "id": "FR-001",
      "title": "Requirement title",
      "description": "Detailed requirement description",
      "priority": "High|Medium|Low",
      "category": "User Interface|Business Logic|Data|Integration|Security"
    }}
  ],
  "non_functional_requirements": [
    {{
      "id": "NFR-001", 
      "title": "Requirement title",
      "description": "Detailed requirement description",
      "category": "Performance|Security|Usability|Reliability|Scalability",
      "metric": "Measurable criteria if applicable"
    }}
  ],
  "stakeholder_roles": [
    {{
      "role": "Role title",
      "responsibilities": ["List of responsibilities"],
      "requirements_impact": "How this role affects requirements"
    }}
  ],
  "technical_considerations": {{
    "architecture_notes": "Key architectural decisions and patterns",
    "integration_points": "External systems and APIs",
    "data_requirements": "Data storage and processing needs",
    "security_requirements": "Security and compliance considerations"
  }},
  "implementation_phases": [
    {{
      "phase": "Phase name",
      "duration_estimate": "Time estimate",
      "deliverables": ["List of deliverables"],
      "dependencies": ["List of dependencies"]
    }}
  ]
}}

Return ONLY the JSON response, no additional text."""

@retry_with_backoff(max_attempts=3, circuit_breaker_name='bedrock')
def call_bedrock_with_retry(prompt: str, filename: str) -> Dict[str, Any]:
    """Call Bedrock with optimized retry logic and caching."""
    
    start_time = time.time()
    
    try:
        # Prepare Bedrock request with optimized parameters
        request_body = {
            "prompt": f"\n\nHuman: {prompt}\n\nAssistant:",
            "max_tokens_to_sample": config.bedrock.max_tokens,
            "temperature": config.bedrock.temperature,
            "top_p": config.bedrock.top_p,
            "stop_sequences": ["\n\nHuman:"]
        }
        
        logger.info("Calling Bedrock API",
                   model_id=config.bedrock.model_id,
                   prompt_length=len(prompt),
                   max_tokens=config.bedrock.max_tokens)
        
        # Call Bedrock with timeout
        response = bedrock_client.invoke_model(
            body=json.dumps(request_body),
            modelId=config.bedrock.model_id,
            accept='application/json',
            contentType='application/json'
        )
        
        # Process response
        response_body = json.loads(response['body'].read())
        completion = response_body.get('completion', '')
        
        duration = time.time() - start_time
        logger.info("Bedrock API call successful",
                   response_length=len(completion),
                   duration_seconds=duration,
                   filename=filename)
        
        # Parse JSON response
        try:
            # Clean up completion text
            json_start = completion.find('{')
            json_end = completion.rfind('}') + 1
            
            if json_start >= 0 and json_end > json_start:
                json_text = completion[json_start:json_end]
                return json.loads(json_text)
            else:
                logger.warning("No valid JSON found in Bedrock response")
                return create_fallback_analysis_from_text(completion, filename)
                
        except json.JSONDecodeError as e:
            logger.error("Failed to parse Bedrock JSON response", error=str(e))
            return create_fallback_analysis_from_text(completion, filename)
            
    except Exception as e:
        duration = time.time() - start_time
        logger.error("Bedrock API call failed",
                    error=str(e),
                    duration_seconds=duration,
                    filename=filename)
        raise ExternalServiceError(f"Bedrock processing failed: {str(e)}")

def create_fallback_analysis(text_content: str, filename: str) -> Dict[str, Any]:
    """Create a basic fallback analysis when Bedrock fails."""
    
    return {
        "document_summary": f"Document analysis for {filename} (fallback mode)",
        "functional_requirements": [
            {
                "id": "FR-001",
                "title": "Document Processing Requirements",
                "description": f"Process and analyze the uploaded document: {filename}",
                "priority": "High",
                "category": "Business Logic"
            }
        ],
        "non_functional_requirements": [
            {
                "id": "NFR-001",
                "title": "System Performance",
                "description": "System should process documents efficiently",
                "category": "Performance",
                "metric": "Process documents within 5 minutes"
            }
        ],
        "stakeholder_roles": [
            {
                "role": "Document Owner",
                "responsibilities": ["Provide clear requirements", "Review analysis results"],
                "requirements_impact": "Defines the scope and detail of analysis needed"
            }
        ],
        "technical_considerations": {
            "architecture_notes": "Basic document processing system",
            "integration_points": "Document upload and analysis APIs",
            "data_requirements": "Document storage and metadata management",
            "security_requirements": "Secure document handling and access control"
        },
        "implementation_phases": [
            {
                "phase": "Initial Analysis",
                "duration_estimate": "1-2 weeks",
                "deliverables": ["Requirements document", "Technical specification"],
                "dependencies": ["Stakeholder approval", "Technical review"]
            }
        ],
        "processing_note": "This analysis was generated in fallback mode due to AI processing limitations.",
        "text_length": len(text_content)
    }

def create_fallback_analysis_from_text(text_response: str, filename: str) -> Dict[str, Any]:
    """Create fallback analysis from partial text response."""
    
    fallback = create_fallback_analysis("", filename)
    fallback["ai_response_text"] = text_response
    fallback["processing_note"] = "Partial AI analysis - JSON parsing failed"
    
    return fallback

@retry_with_backoff(max_attempts=3)
def update_processing_stage_optimized(request_id: str, stage: str, additional_data: Dict[str, Any] = None) -> None:
    """Update processing stage with optimized DynamoDB operations."""
    
    try:
        table = dynamodb.Table(HISTORY_TABLE)
        
        # Prepare update expression
        update_expression = "SET processingStage = :stage, lastUpdated = :timestamp"
        expression_values = {
            ':stage': stage,
            ':timestamp': datetime.now(timezone.utc).isoformat()
        }
        
        # Add additional data if provided
        if additional_data:
            for key, value in additional_data.items():
                update_expression += f", {key} = :{key}"
                expression_values[f":{key}"] = value
        
        # Update with condition to ensure record exists
        table.update_item(
            Key={'requestId': request_id, 'timestamp': request_id},
            UpdateExpression=update_expression,
            ExpressionAttributeValues=expression_values,
            ConditionExpression='attribute_exists(requestId)'
        )
        
        logger.info("Processing stage updated",
                   request_id=request_id,
                   stage=stage,
                   additional_fields=list(additional_data.keys()) if additional_data else [])
                   
    except Exception as e:
        logger.error("Failed to update processing stage",
                    request_id=request_id,
                    stage=stage,
                    error=str(e))
        raise DatabaseError(f"Failed to update processing stage: {str(e)}")

@retry_with_backoff(max_attempts=3)
def store_results_optimized(request_id: str, ai_analysis: Dict[str, Any], db_record: Dict[str, Any]) -> None:
    """Store AI analysis results with optimized DynamoDB operations."""
    
    try:
        table = dynamodb.Table(HISTORY_TABLE)
        
        # Prepare the analysis data
        analysis_data = {
            'aiAnalysis': ai_analysis,
            'processedAt': datetime.now(timezone.utc).isoformat(),
            'analysisMetadata': {
                'functional_requirements_count': len(ai_analysis.get('functional_requirements', [])),
                'non_functional_requirements_count': len(ai_analysis.get('non_functional_requirements', [])),
                'stakeholder_roles_count': len(ai_analysis.get('stakeholder_roles', [])),
                'implementation_phases_count': len(ai_analysis.get('implementation_phases', [])),
                'analysis_size_bytes': len(json.dumps(ai_analysis)),
                'cache_enabled': True
            }
        }
        
        # Update the main record
        table.update_item(
            Key={'requestId': request_id, 'timestamp': db_record['timestamp']},
            UpdateExpression="""
                SET aiAnalysis = :analysis,
                    processedAt = :processed_at,
                    analysisMetadata = :metadata,
                    status = :status
            """,
            ExpressionAttributeValues={
                ':analysis': ai_analysis,
                ':processed_at': analysis_data['processedAt'],
                ':metadata': analysis_data['analysisMetadata'],
                ':status': 'processed'
            }
        )
        
        logger.info("AI analysis results stored",
                   request_id=request_id,
                   analysis_size_bytes=analysis_data['analysisMetadata']['analysis_size_bytes'])
                   
    except Exception as e:
        logger.error("Failed to store AI analysis results",
                    request_id=request_id,
                    error=str(e))
        raise DatabaseError(f"Failed to store analysis results: {str(e)}")

def handle_health_check() -> Dict[str, Any]:
    """Handle health check requests with cache statistics."""
    
    try:
        # Create health checker with all dependencies
        health_checker = HealthChecker('process')
        
        # Add service-specific health checks
        health_checker.add_dependency_check('s3', AWSHealthChecks.create_s3_check(DOCUMENT_BUCKET))
        health_checker.add_dependency_check('dynamodb', AWSHealthChecks.create_dynamodb_check(HISTORY_TABLE))
        health_checker.add_dependency_check('bedrock', AWSHealthChecks.create_bedrock_check(config.bedrock.model_id))
        
        # Perform health check
        health_result = health_checker.check_health()
        
        # Add circuit breaker status
        circuit_breaker_status = {
            'bedrock_circuit_breaker': bedrock_circuit_breaker.get_state()
        }
        
        # Get cache statistics
        cache_stats = cache_manager.get_stats()
        
        return {
            'statusCode': 200 if health_result.status.value == 'healthy' else 503,
            'body': json.dumps({
                'health': health_result.to_dict(),
                'circuit_breakers': circuit_breaker_status,
                'cache_stats': cache_stats,
                'function_config': {
                    'memory_mb': os.environ.get('AWS_LAMBDA_FUNCTION_MEMORY_SIZE'),
                    'timeout_seconds': os.environ.get('AWS_LAMBDA_FUNCTION_TIMEOUT'),
                    'environment': config.environment,
                    'caching_enabled': True
                }
            })
        }
        
    except Exception as e:
        logger.error("Health check failed", error=str(e))
        return {
            'statusCode': 503,
            'body': json.dumps({
                'health': 'unhealthy',
                'error': str(e),
                'caching_enabled': True
            })
        }